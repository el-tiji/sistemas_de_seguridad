from fastapi import APIRouter, Depends, HTTPException, Form
from sqlalchemy.orm import Session
from sqlalchemy import text
from fastapi.responses import FileResponse
from docx import Document
import os

from src.backend.config.db import get_db

router = APIRouter(
    prefix="/api/soa",
    tags=["SoA"]
)

# =========================================
# GENERAR SOA
# =========================================
@router.post("/generate")
def generate_soa(
    organizacion_id: int = Form(...),
    responsable: str = Form(...),
    descripcion: str = Form(...),
    db: Session = Depends(get_db)
):

    try:

        # =========================================
        # OBTENER CONTROLES
        # =========================================
        controls = db.execute(
            text("""
                SELECT id
                FROM control
            """)
        ).fetchall()

        if not controls:
            raise HTTPException(
                status_code=400,
                detail="No existen controles registrados"
            )

        # =========================================
        # OBTENER ÚLTIMA VERSION
        # =========================================
        last_version = db.execute(
            text("""
                SELECT MAX(version)
                FROM soa
            """)
        ).scalar()

        if last_version is None:
            new_version = 1
        else:
            new_version = last_version + 1

        # =========================================
        # CREAR SOA
        # =========================================
        db.execute(
            text("""
                INSERT INTO soa (
                    organizacion_id,
                    version,
                    responsable,
                    descripcion,
                    estado
                )
                VALUES (
                    :organizacion_id,
                    :version,
                    :responsable,
                    :descripcion,
                    :estado
                )
            """),
            {
                "organizacion_id": organizacion_id,
                "version": new_version,
                "responsable": responsable,
                "descripcion": descripcion,
                "estado": "ACTIVO"
            }
        )

        db.commit()

        # =========================================
        # OBTENER ID DEL SOA
        # =========================================
        soa_id = db.execute(
            text("SELECT LAST_INSERT_ID()")
        ).scalar()

        # =========================================
        # INSERTAR CONTROLES
        # =========================================
        for control in controls:

            db.execute(
                text("""
                    INSERT INTO soa_control (
                        soa_id,
                        control_id,
                        aplica,
                        justificacion_inclusion,
                        justificacion_exclusion,
                        estado_implementacion
                    )
                    VALUES (
                        :soa_id,
                        :control_id,
                        :aplica,
                        :justificacion_inclusion,
                        :justificacion_exclusion,
                        :estado_implementacion
                    )
                """),
                {
                    "soa_id": soa_id,
                    "control_id": control[0],
                    "aplica": True,
                    "justificacion_inclusion": "Control requerido para SGSI",
                    "justificacion_exclusion": None,
                    "estado_implementacion": "PENDIENTE"
                }
            )

        db.commit()

        return {
            "message": "SoA generada correctamente",
            "soa_id": soa_id,
            "version": new_version,
            "total_controles": len(controls)
        }

    except HTTPException:
        raise

    except Exception as e:

        db.rollback()

        raise HTTPException(
            status_code=500,
            detail=f"Error al generar SoA: {str(e)}"
        )

# =========================================
# LISTAR SOA
# =========================================
@router.get("/list-soa")
def listar_soa(
    db: Session = Depends(get_db)
):

    try:

        result = db.execute(text("""
            SELECT
                s.id,
                s.version,
                s.fecha,
                s.responsable,
                s.estado,
                o.nombre
            FROM soa s
            JOIN organizacion o
                ON s.organizacion_id = o.id
            ORDER BY s.id DESC
        """))

        rows = result.fetchall()

        return [
            {
                "id": r[0],
                "version": r[1],
                "fecha": str(r[2]),
                "responsable": r[3],
                "estado": r[4],
                "organizacion": r[5]
            }
            for r in rows
        ]

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=f"Error al listar SoA: {str(e)}"
        )

# =========================================
# DETALLE DEL SOA
# =========================================
@router.get("/detail/{soa_id}")
def detalle_soa(
    soa_id: int,
    db: Session = Depends(get_db)
):

    try:

        result = db.execute(
            text("""
                SELECT
                    c.codigo,
                    c.descripcion,
                    sc.aplica,
                    sc.justificacion_inclusion,
                    sc.justificacion_exclusion,
                    sc.estado_implementacion
                FROM soa_control sc
                JOIN control c
                    ON sc.control_id = c.id
                WHERE sc.soa_id = :soa_id
            """),
            {"soa_id": soa_id}
        )

        rows = result.fetchall()

        return [
            {
                "codigo": r[0],
                "descripcion": r[1],
                "aplica": r[2],
                "justificacion_inclusion": r[3],
                "justificacion_exclusion": r[4],
                "estado_implementacion": r[5]
            }
            for r in rows
        ]

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener detalle SoA: {str(e)}"
        )



# =========================================
# GENERAR DOCUMENTO WORD DEL SOA
# =========================================
@router.get("/generate-doc/{soa_id}")
def generar_documento_soa(
    soa_id: int,
    db: Session = Depends(get_db)
):

    try:

        # =========================================
        # OBTENER SOA
        # =========================================
        soa = db.execute(text("""
            SELECT
                s.id,
                s.version,
                s.fecha,
                s.responsable,
                s.descripcion,
                s.estado,
                o.nombre
            FROM soa s
            JOIN organizacion o
                ON s.organizacion_id = o.id
            WHERE s.id = :soa_id
        """), {
            "soa_id": soa_id
        }).fetchone()

        if not soa:
            raise HTTPException(
                status_code=404,
                detail="SoA no encontrada"
            )

        # =========================================
        # OBTENER CONTROLES
        # =========================================
        controles = db.execute(text("""
            SELECT
                c.codigo,
                c.descripcion,
                sc.aplica,
                sc.justificacion_inclusion,
                sc.estado_implementacion
            FROM soa_control sc
            JOIN control c
                ON sc.control_id = c.id
            WHERE sc.soa_id = :soa_id
        """), {
            "soa_id": soa_id
        }).fetchall()

        # =========================================
        # CREAR DOCUMENTO
        # =========================================
        doc = Document()

        doc.add_heading(
            'Declaración de Aplicabilidad (SoA)',
            level=1
        )

        doc.add_paragraph(f"Organización: {soa[6]}")
        doc.add_paragraph(f"Versión: {soa[1]}")
        doc.add_paragraph(f"Fecha: {soa[2]}")
        doc.add_paragraph(f"Responsable: {soa[3]}")
        doc.add_paragraph(f"Estado: {soa[5]}")
        doc.add_paragraph(f"Descripción: {soa[4]}")

        doc.add_heading('Controles', level=2)

        # =========================================
        # TABLA
        # =========================================
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells

        hdr[0].text = 'Código'
        hdr[1].text = 'Descripción'
        hdr[2].text = 'Aplica'
        hdr[3].text = 'Justificación'
        hdr[4].text = 'Estado'

        for control in controles:

            row = table.add_row().cells

            row[0].text = str(control[0])
            row[1].text = str(control[1])
            row[2].text = 'Sí' if control[2] else 'No'
            row[3].text = str(control[3])
            row[4].text = str(control[4])

        # =========================================
        # GUARDAR ARCHIVO
        # =========================================
        output_dir = "generated_docs"

        os.makedirs(output_dir, exist_ok=True)

        file_path = f"{output_dir}/soa_{soa_id}.docx"

        doc.save(file_path)

        # =========================================
        # RETORNAR ARCHIVO
        # =========================================
        return FileResponse(
            path=file_path,
            filename=f"soa_{soa_id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except HTTPException:
        raise

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=f"Error al generar documento: {str(e)}"
        )